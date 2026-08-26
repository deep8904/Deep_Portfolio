#!/usr/bin/env python3
"""Build the tagged, ATS-friendly portfolio resume PDF from verified content."""

from pathlib import Path
import json
import os
import shutil
import subprocess
import tempfile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCX_PATH = ROOT / "scripts" / "resume" / "Deep-Chadamiya-Resume.docx"
PDF_DIR = ROOT / "public" / "resume"

INK = RGBColor(30, 29, 27)
MUTED = RGBColor(91, 87, 81)
ACCENT = RGBColor(139, 91, 57)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "8B5B39")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([props, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_rule(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "D7D0C8")
    borders.append(bottom)
    p_pr.append(borders)


def add_section_heading(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    add_rule(p)
    return p


def add_role(doc, title, org, location, dates, bullets):
    p = doc.add_paragraph(style="Resume Role")
    run = p.add_run(f"{title} | {org}")
    run.bold = True
    meta = doc.add_paragraph(f"{location} | {dates}", style="Resume Meta")
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")


def add_project(doc, title, descriptor, stack, bullets):
    p = doc.add_paragraph(style="Resume Role")
    run = p.add_run(f"{title} | {descriptor}")
    run.bold = True
    doc.add_paragraph(stack, style="Resume Meta")
    for bullet in bullets:
        doc.add_paragraph(bullet, style="List Bullet")


def build_docx():
    doc = Document()
    props = doc.core_properties
    props.title = "Deepkumar Chadamiya Resume"
    props.subject = "Product Designer, Design Engineer, and Front-End Developer"
    props.author = "Deepkumar Chadamiya"
    props.keywords = "product design, design engineering, frontend development, UX, React, Next.js"

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(2.4)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title_style = styles["Title"]
    title_style.font.name = "Arial"
    title_style.font.size = Pt(23)
    title_style.font.bold = True
    title_style.font.color.rgb = INK
    title_style.paragraph_format.space_after = Pt(1)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h1 = styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(10.8)
    h1.font.bold = True
    h1.font.color.rgb = ACCENT
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(3)
    h1.paragraph_format.keep_with_next = True

    role_style = styles.add_style("Resume Role", WD_STYLE_TYPE.PARAGRAPH)
    role_style.font.name = "Arial"
    role_style.font.size = Pt(9.6)
    role_style.font.color.rgb = INK
    role_style.paragraph_format.space_before = Pt(3.5)
    role_style.paragraph_format.space_after = Pt(0)
    role_style.paragraph_format.keep_with_next = True

    meta_style = styles.add_style("Resume Meta", WD_STYLE_TYPE.PARAGRAPH)
    meta_style.font.name = "Arial"
    meta_style.font.size = Pt(8.4)
    meta_style.font.color.rgb = MUTED
    meta_style.paragraph_format.space_after = Pt(1.5)
    meta_style.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet.font.size = Pt(8.8)
    bullet.font.color.rgb = INK
    bullet.paragraph_format.left_indent = Inches(0.18)
    bullet.paragraph_format.first_line_indent = Inches(-0.12)
    bullet.paragraph_format.space_after = Pt(1.2)

    doc.add_paragraph("Deepkumar Chadamiya", style="Title")
    subtitle = doc.add_paragraph("Product Designer | Design Engineer | Front-End Developer")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    subtitle.runs[0].font.size = Pt(10.5)
    subtitle.runs[0].font.bold = True

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(5)
    contact.add_run("Tempe, AZ | +1 (480) 572-6950 | ")
    add_hyperlink(contact, "deeppatel8904@gmail.com", "mailto:deeppatel8904@gmail.com")
    contact.add_run(" | ")
    add_hyperlink(contact, "linkedin.com/in/deepchadamiya", "https://www.linkedin.com/in/deepchadamiya")
    contact.add_run(" | ")
    add_hyperlink(contact, "github.com/deep8904", "https://github.com/deep8904")

    add_section_heading(doc, "SUMMARY")
    doc.add_paragraph(
        "Product Designer and Design Engineer building digital products across professional, academic, and independent work. "
        "Experience spans product design, UI/UX, frontend development, design systems, game technology, data visualization, "
        "and AI-powered products. Comfortable taking ideas from research and user flows through prototypes, responsive interfaces, "
        "production code, APIs, testing, and deployment."
    )

    add_section_heading(doc, "SKILLS")
    skill_rows = [
        ("Product and UX", "Product Design, UI/UX Design, Interaction Design, Information Architecture, User Flows, Wireframing, Prototyping, User Research, Usability Testing, Accessibility, WCAG"),
        ("Design systems and tools", "Figma, Framer, Design Systems, Component Libraries, Auto Layout, Variables, Design Tokens, Responsive Design, Data Visualization"),
        ("Frontend", "React, Next.js, TypeScript, JavaScript, HTML5, CSS3, Tailwind CSS, Angular"),
        ("Backend, data, and delivery", "Node.js, Java, REST APIs, Supabase, PostgreSQL, Git, GitHub Actions, Vercel"),
        ("AI and prototyping", "LLM Integration, Gemini API, Claude Code, Structured Outputs, AI-Assisted Prototyping, Workflow Automation, Human-in-the-Loop Workflows, Playtesting"),
    ]
    for label, values in skill_rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1.3)
        r = p.add_run(f"{label}: ")
        r.bold = True
        p.add_run(values)

    add_section_heading(doc, "EXPERIENCE")
    add_role(doc, "Software and Game Development", "Endless Games Studio", "Mesa, AZ", "Aug 2026 - Present", [
        "Support software and game development across lab programs through prototyping, coding support, playtesting, feature testing, bug documentation, and technical troubleshooting.",
        "Assist with development environments, builds, demo stations, workshops, and events while helping students and participants solve technical problems during hands-on projects.",
        "Collaborate on game development and special projects, including Endstar platform support, while identifying technical issues and communicating project risks to the team.",
    ])
    add_role(doc, "Design Technology Teaching Assistant / Database Assistant", "Arizona State University", "Tempe, AZ", "Oct 2024 - Dec 2025", [
        "Assisted design students with digital tools, prototyping, visual communication, portfolio development, and technical implementation.",
        "Designed Airtable systems with linked records, structured fields, views, and automated workflows, reducing manual processing time by 60% across 3 academic datasets.",
        "Built 6 Power BI and Tableau dashboards across 5+ programs, turning enrollment, academic standing, and graduation data into clear reports for faculty and staff.",
    ])
    add_role(doc, "Software Developer / Full-Stack Developer", "Tibicle LLP", "Ahmedabad, India", "Dec 2022 - Nov 2023", [
        "Worked across product, design, and engineering on SaaS and EdTech products using Angular, TypeScript, Java, REST APIs, and relational databases.",
        "Built reusable UI components, forms, dashboards, navigation, and role-based workflows for platforms supporting 500+ users.",
        "Improved frontend architecture, state management, and API performance, reducing page load times by 40%; rebuilt third-party integrations and reduced integration defects by 60%.",
    ])

    doc.add_page_break()
    add_section_heading(doc, "SELECTED PROJECTS")
    add_project(doc, "CreatorFlow", "AI Workspace for Creators", "Product Design | React | Next.js | TypeScript | Supabase | AI", [
        "Designed and built a full-stack platform for sponsorship management, content planning, analytics, AI-assisted repurposing, automation, and team collaboration.",
        "Defined information architecture, user flows, responsive interfaces, reusable UI patterns, and role-based experiences across 5 user roles, then implemented them in React and Next.js.",
        "Built authentication, PostgreSQL row-level security, permissions, Edge Functions, and Gmail and YouTube API integrations.",
    ])
    add_project(doc, "C.A.R.E. for Horses", "UX Research and Website Redesign", "UX Research | UI/UX Design | Prototyping | Accessibility", [
        "Conducted surveys, heuristic evaluation, competitive analysis, personas, and journey mapping to identify navigation, usability, accessibility, and content problems.",
        "Used findings from 16 survey responses to redesign information architecture, user flows, page hierarchy, and interaction patterns around clearer user tasks.",
    ])
    add_project(doc, "ACM", "AI Content Publishing System", "TypeScript | Gemini API | PostgreSQL | GitHub Actions | Vercel", [
        "Designed and built an AI workflow for trend discovery, source-backed research, content generation, editorial review, human approval, publishing, and deployment verification.",
        "Built structured LLM workflows with Gemini API, Zod validation, PostgreSQL state, retry handling, approval gates, and resumable automation.",
    ])

    add_section_heading(doc, "EDUCATION")
    add_role(doc, "Master of Science in Information Technology", "Arizona State University", "Tempe, AZ", "Dec 2025", ["GPA: 4.0/4.0"])
    add_role(doc, "Bachelor of Engineering in Computer Engineering", "Vidush Somany Institute of Technology and Research", "India", "Apr 2023", ["GPA: 8.95/10.0"])

    for sec in doc.sections:
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_before = Pt(3)
        run = footer.add_run("Deepkumar Chadamiya | Resume")
        run.font.name = "Arial"
        run.font.size = Pt(7.5)
        run.font.color.rgb = MUTED

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def export_pdf():
    soffice = shutil.which("soffice")
    if not soffice:
        raise SystemExit("LibreOffice (soffice) is required to export the tagged PDF")
    filter_options = {
        "UseTaggedPDF": {"type": "boolean", "value": "true"},
        "PDFUACompliance": {"type": "boolean", "value": "true"},
        "ExportBookmarks": {"type": "boolean", "value": "true"},
        "ExportLinksRelativeFsys": {"type": "boolean", "value": "false"},
    }
    target = PDF_DIR / "Deep-Chadamiya-Resume.pdf"
    with tempfile.TemporaryDirectory(prefix="deep-resume-") as temp_dir:
        profile = Path(temp_dir) / "libreoffice-profile"
        output_dir = Path(temp_dir) / "output"
        output_dir.mkdir()
        env = os.environ.copy()
        env["XDG_CACHE_HOME"] = str(Path(temp_dir) / "cache")
        subprocess.run(
            [
                soffice,
                f"-env:UserInstallation={profile.as_uri()}",
                "--headless",
                "--convert-to",
                f"pdf:writer_pdf_Export:{json.dumps(filter_options, separators=(',', ':'))}",
                "--outdir",
                str(output_dir),
                str(DOCX_PATH),
            ],
            check=True,
            env=env,
        )
        generated = output_dir / "Deep-Chadamiya-Resume.pdf"
        if not generated.exists():
            raise SystemExit("LibreOffice did not produce the expected PDF")
        shutil.copy2(generated, target)


if __name__ == "__main__":
    build_docx()
    export_pdf()
    print(PDF_DIR / "Deep-Chadamiya-Resume.pdf")
